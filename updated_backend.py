#!/usr/bin/env python3
"""
NOVA ULTRA PROFESSIONAL FASTAPI BACKEND
Complete integration of all enhanced_cli.py functionality with REST API endpoints
"""

import asyncio
import os
import sys
import json
import time
import threading
import sqlite3
import logging
import hashlib
import re
import requests
import random
import pickle
import base64
import subprocess
import webbrowser
import geocoder
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Union, Tuple
from collections import defaultdict, deque, Counter
from pathlib import Path
import numpy as np

# FastAPI imports
from fastapi import FastAPI, HTTPException, Request, Depends
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
import uvicorn

# Multi-folder import solution
project_root = os.path.dirname(os.path.abspath(__file__))
folders_to_add = [
    'src',
    os.path.join('src', 'memory'),
    os.path.join('src', 'unique_features'),
    os.path.join('src', 'agents'),
]

for folder in folders_to_add:
    folder_path = os.path.join(project_root, folder)
    if os.path.exists(folder_path) and folder_path not in sys.path:
        sys.path.insert(0, folder_path)

from dotenv import load_dotenv
load_dotenv()

# Rich UI imports (for logging)
try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
    from rich.table import Table
    from rich.markdown import Markdown
    from rich.text import Text
    console = Console()
    RICH_AVAILABLE = True
except ImportError:
    RICH_AVAILABLE = False

# Voice processing imports (Basic + Azure)
try:
    import speech_recognition as sr
    import pyttsx3
    VOICE_AVAILABLE = True
except ImportError:
    VOICE_AVAILABLE = False

# Azure Voice imports (PREMIUM)
try:
    import azure.cognitiveservices.speech as speechsdk
    AZURE_VOICE_AVAILABLE = True
except ImportError:
    AZURE_VOICE_AVAILABLE = False

# File processing imports
try:
    from PIL import Image, ImageEnhance, ImageFilter
    import PyPDF2
    import docx
    import openpyxl
    import matplotlib.pyplot as plt
    import seaborn as sns
    import pandas as pd
    import cv2
    FILE_PROCESSING_AVAILABLE = True
except ImportError:
    FILE_PROCESSING_AVAILABLE = False

# Web scraping and search imports (PREMIUM)
try:
    from bs4 import BeautifulSoup
    import feedparser
    WEB_SEARCH_AVAILABLE = True
except ImportError:
    WEB_SEARCH_AVAILABLE = False

# Image generation imports (PREMIUM)
try:
    import replicate
    IMAGE_GENERATION_AVAILABLE = True
except ImportError:
    IMAGE_GENERATION_AVAILABLE = False

# GitHub Integration imports
try:
    import chromadb
    # Try new langchain imports first with fallback
    try:
        from langchain_community.document_loaders import UnstructuredFileLoader
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain_community.vectorstores import Chroma
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except ImportError:
        # Fallback to safe text loader if unstructured fails
        try:
            from langchain_community.document_loaders import TextLoader as UnstructuredFileLoader
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from langchain_community.vectorstores import Chroma
            from langchain_community.embeddings import HuggingFaceEmbeddings
        except ImportError:
            # Final fallback to old imports
            from langchain.document_loaders import TextLoader as UnstructuredFileLoader
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from langchain.vectorstores import Chroma
            from langchain.embeddings import HuggingFaceEmbeddings
    GITHUB_INTEGRATION = True
except ImportError as e:
    GITHUB_INTEGRATION = False

# Professional Agents Import
try:
    from agents.coding_agent import ProLevelCodingExpert
    from agents.career_coach import ProfessionalCareerCoach
    from agents.business_consultant import SmartBusinessConsultant
    from agents.medical_advisor import SimpleMedicalAdvisor
    from agents.emotional_counselor import SimpleEmotionalCounselor
    from agents.techincal_architect import TechnicalArchitect
    PROFESSIONAL_AGENTS_LOADED = True
except ImportError as e:
    PROFESSIONAL_AGENTS_LOADED = False

# Advanced Systems Import
try:
    from memory.sharp_memory import SharpMemorySystem
    from unique_features.smart_orchestrator import IntelligentAPIOrchestrator
    from unique_features.api_drift_detector import APIPerformanceDrifter
    ADVANCED_SYSTEMS = True
except ImportError as e:
    ADVANCED_SYSTEMS = False

# GitHub Repo Analysis Import
try:
    from agents.ingest import main as ingest_repo, process_and_store_documents
    from agents.qa_engine import create_qa_engine, EnhancedQAEngine
    GITHUB_INTEGRATION = GITHUB_INTEGRATION and True
except ImportError as e:
    GITHUB_INTEGRATION = False
    ingest_repo = None
    create_qa_engine = None

# ========== PYDANTIC MODELS ==========
class UserInput(BaseModel):
    text: str
    user_id: str = "default_user"
    session_id: str = None
    language: str = "en-IN"
    voice_mode: bool = False

class VoiceInput(BaseModel):
    audio_data: str  # Base64 encoded
    language: str = "en-IN"
    user_id: str = "default_user"

class FileProcessRequest(BaseModel):
    file_path: str
    user_id: str = "default_user"

class GitHubAnalysisRequest(BaseModel):
    repo_url: str
    user_id: str = "default_user"

class ImageGenerationRequest(BaseModel):
    prompt: str
    size: str = "1024x1024"
    user_id: str = "default_user"

class WeatherRequest(BaseModel):
    location: str = None
    user_id: str = "default_user"

class SearchRequest(BaseModel):
    query: str
    max_results: int = 5
    user_id: str = "default_user"

class CryptoRequest(BaseModel):
    symbols: List[str] = None
    user_id: str = "default_user"

class MemoryRequest(BaseModel):
    user_id: str
    limit: int = 15

# ========== FASTAPI APP SETUP ==========
app = FastAPI(
    title="NOVA Ultra Professional AI System",
    description="Ultra Advanced AI Assistant with ALL Premium Features",
    version="2.0.0",
    docs_url="/api/docs",
    redoc_url="/api/redoc"
)

# CORS Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# ========== CORE SYSTEMS ==========
class Colors:
    RESET = '\033[0m'
    BOLD = '\033[1m'
    RED = '\033[91m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    PURPLE = '\033[95m'
    MAGENTA = '\033[95m'
    WHITE = '\033[97m'
    ORANGE = '\033[38;5;208m'

class UltraHybridMemorySystem:
    """Ultra Advanced Hybrid Memory with ALL previous features"""
    
    def __init__(self, db_path="nova_ultra_professional_memory.db"):
        if not os.path.isabs(db_path):
            self.db_path = os.path.join(os.getcwd(), db_path)
        else:
            self.db_path = db_path
        self.setup_database()
        
        self.conversation_context = deque(maxlen=100)
        self.user_profile = {}
        self.emotional_state = "neutral"
        self.learning_patterns = defaultdict(list)
        self.personality_insights = {}
        self.user_preferences = {}
        self.conversation_history = []
        self.short_term_memory = deque(maxlen=200)
        self.working_memory = {}
        self.conversation_threads = {}
        self.context_memory = {}
        self.voice_memory = deque(maxlen=50)
        self.file_memory = {}
        self.search_memory = deque(maxlen=30)
        self.image_memory = deque(maxlen=20)
        
        self.setup_semantic_memory()

    def setup_database(self):
        """Setup ultra comprehensive database schema"""
        try:
            db_dir = os.path.dirname(self.db_path)
            if db_dir and not os.path.exists(db_dir):
                os.makedirs(db_dir, exist_ok=True)
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS conversations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    session_id TEXT,
                    user_input TEXT,
                    bot_response TEXT,
                    agent_type TEXT,
                    language TEXT,
                    emotion TEXT,
                    confidence REAL,
                    timestamp DATETIME,
                    feedback INTEGER DEFAULT 0,
                    context_summary TEXT,
                    learned_facts TEXT,
                    satisfaction_rating INTEGER,
                    conversation_thread_id TEXT,
                    intent_detected TEXT,
                    response_time REAL,
                    voice_used BOOLEAN DEFAULT 0,
                    location TEXT,
                    weather_context TEXT,
                    search_queries TEXT
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS user_profiles (
                    user_id TEXT PRIMARY KEY,
                    name TEXT,
                    career_goals TEXT,
                    current_role TEXT,
                    experience_years INTEGER,
                    skills TEXT,
                    preferences TEXT,
                    communication_style TEXT,
                    emotional_patterns TEXT,
                    conversation_patterns TEXT,
                    expertise_level TEXT,
                    topics_of_interest TEXT,
                    last_updated DATETIME,
                    total_conversations INTEGER DEFAULT 0,
                    preferred_voice TEXT,
                    location TEXT,
                    timezone TEXT,
                    personality_type TEXT,
                    learning_style TEXT
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS github_repos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    repo_url TEXT UNIQUE,
                    repo_name TEXT,
                    analysis_date DATETIME,
                    file_count INTEGER,
                    languages_detected TEXT,
                    issues_found TEXT,
                    suggestions TEXT,
                    vector_db_path TEXT
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS voice_interactions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    voice_input TEXT,
                    voice_response TEXT,
                    language_detected TEXT,
                    emotion_detected TEXT,
                    voice_engine TEXT,
                    timestamp DATETIME
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS file_processing (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    file_path TEXT,
                    file_type TEXT,
                    processing_result TEXT,
                    timestamp DATETIME,
                    success BOOLEAN
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS search_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    search_query TEXT,
                    search_type TEXT,
                    results_count INTEGER,
                    timestamp DATETIME
                )
                ''')
                
                cursor.execute('''
                CREATE TABLE IF NOT EXISTS weather_queries (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT,
                    location TEXT,
                    weather_data TEXT,
                    timestamp DATETIME
                )
                ''')
                
                conn.commit()
            
        except Exception as e:
            print(f"Database setup error: {e}")

    def setup_semantic_memory(self):
        """Setup semantic memory for technical queries"""
        try:
            if ADVANCED_SYSTEMS:
                self.semantic_memory = SharpMemorySystem()
            else:
                self.semantic_memory = None
        except Exception as e:
            print(f"Semantic memory setup error: {e}")
            self.semantic_memory = None

    async def remember_conversation(self, user_id: str, session_id: str,
                                  user_input: str, bot_response: str,
                                  agent_type: str, language: str,
                                  emotion: str, confidence: float,
                                  intent: str = None, response_time: float = 0.0,
                                  voice_used: bool = False, location: str = None,
                                  weather_context: str = None, search_queries: str = None):
        """Ultra enhanced conversation memory storage"""
        try:
            learned_facts = self.extract_learning_points(user_input, bot_response)
            context_summary = self.generate_context_summary()
            
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO conversations 
                (user_id, session_id, user_input, bot_response, agent_type,
                 language, emotion, confidence, timestamp, context_summary,
                 learned_facts, conversation_thread_id, intent_detected, response_time,
                 voice_used, location, weather_context, search_queries)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (user_id, session_id, user_input, bot_response, agent_type,
                      language, emotion, confidence, datetime.now(), context_summary,
                      learned_facts, self.generate_thread_id(), intent, response_time,
                      voice_used, location, weather_context, search_queries))
                conn.commit()
            
            self.conversation_context.append({
                'user': user_input,
                'bot': bot_response,
                'emotion': emotion,
                'agent': agent_type,
                'timestamp': datetime.now(),
                'voice_used': voice_used,
                'location': location
            })
            
            memory_entry = {
                'user_id': user_id,
                'session_id': session_id,
                'timestamp': datetime.now().isoformat(),
                'user_input': user_input,
                'ai_response': bot_response,
                'agent_used': agent_type,
                'emotion': emotion,
                'intent': intent,
                'voice_used': voice_used
            }
            self.short_term_memory.append(memory_entry)
            
            if self.semantic_memory and agent_type in ['coding', 'business', 'technical']:
                try:
                    await self.semantic_memory.remember_conversation_advanced(
                        user_input, bot_response,
                        {'agent_used': agent_type, 'emotion': emotion},
                        user_id, session_id
                    )
                except Exception as e:
                    print(f"Semantic memory storage error: {e}")
            
        except Exception as e:
            print(f"Memory storage error: {e}")

    def remember_voice_interaction(self, user_id: str, voice_input: str,
                                 voice_response: str, language: str,
                                 emotion: str, voice_engine: str):
        """Remember voice interactions"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO voice_interactions 
                (user_id, voice_input, voice_response, language_detected,
                 emotion_detected, voice_engine, timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (user_id, voice_input, voice_response, language,
                      emotion, voice_engine, datetime.now()))
                conn.commit()
            
            self.voice_memory.append({
                'input': voice_input,
                'response': voice_response,
                'language': language,
                'emotion': emotion,
                'engine': voice_engine,
                'timestamp': datetime.now()
            })
            
        except Exception as e:
            print(f"Voice memory error: {e}")

    def remember_file_processing(self, user_id: str, file_path: str,
                               file_type: str, result: str, success: bool):
        """Remember file processing"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO file_processing 
                (user_id, file_path, file_type, processing_result, timestamp, success)
                VALUES (?, ?, ?, ?, ?, ?)
                ''', (user_id, file_path, file_type, result, datetime.now(), success))
                conn.commit()
            
            self.file_memory[file_path] = {
                'type': file_type,
                'result': result,
                'success': success,
                'timestamp': datetime.now()
            }
            
        except Exception as e:
            print(f"File memory error: {e}")

    def remember_search_query(self, user_id: str, query: str,
                            search_type: str, results_count: int):
        """Remember search queries"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                INSERT INTO search_history 
                (user_id, search_query, search_type, results_count, timestamp)
                VALUES (?, ?, ?, ?, ?)
                ''', (user_id, query, search_type, results_count, datetime.now()))
                conn.commit()
            
            self.search_memory.append({
                'query': query,
                'type': search_type,
                'count': results_count,
                'timestamp': datetime.now()
            })
            
        except Exception as e:
            print(f"Search memory error: {e}")

    def get_relevant_context(self, user_input: str, user_id: str, limit: int = 15) -> str:
        """Get ultra comprehensive relevant context"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                SELECT user_input, bot_response, emotion, learned_facts, agent_type,
                       voice_used, location, weather_context
                FROM conversations 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT ?
                ''', (user_id, limit))
                conversations = cursor.fetchall()
            
            if not conversations:
                return ""
            
            context = "Previous conversation context:\n"
            for conv in conversations:
                context += f"[{conv[4].upper()}] User ({conv[2]}): {conv[0][:80]}...\n"
                context += f"NOVA: {conv[1][:80]}...\n"
                if conv[3]:
                    context += f"Learned: {conv[3]}\n"
                if conv[5]:
                    context += f"[VOICE MODE]\n"
                if conv[6]:
                    context += f"Location: {conv[6]}\n"
                if conv[7]:
                    context += f"Weather: {conv[7]}\n"
                context += "---\n"
            
            return context
            
        except Exception as e:
            print(f"Context retrieval error: {e}")
            return ""

    def extract_learning_points(self, user_input: str, bot_response: str) -> str:
        """Extract learning points from conversation"""
        learning_keywords = [
            "my name is", "i am", "i work", "i like", "i don't like",
            "my preference", "remember that", "important", "my goal",
            "my project", "my problem", "i need help with", "my role",
            "my company", "my experience", "my skills", "career goal",
            "i live in", "my location", "my city", "my country",
            "i prefer", "i want", "i need", "i use", "my favorite"
        ]
        
        learned = []
        user_lower = user_input.lower()
        
        for keyword in learning_keywords:
            if keyword in user_lower:
                sentences = user_input.split('.')
                for sentence in sentences:
                    if keyword in sentence.lower():
                        learned.append(sentence.strip())
        
        return "; ".join(learned)

    def generate_context_summary(self) -> str:
        """Generate ultra context summary from recent conversations"""
        if not self.conversation_context:
            return ""
        
        recent_topics = []
        emotions = []
        agents = []
        voice_usage = []
        locations = []
        
        for conv in list(self.conversation_context)[-10:]:
            recent_topics.append(conv['user'][:50])
            emotions.append(conv['emotion'])
            agents.append(conv['agent'])
            if conv.get('voice_used'):
                voice_usage.append(True)
            if conv.get('location'):
                locations.append(conv['location'])
        
        dominant_emotion = max(set(emotions), key=emotions.count) if emotions else "neutral"
        most_used_agent = max(set(agents), key=agents.count) if agents else "general"
        voice_percentage = (len(voice_usage) / len(emotions)) * 100 if emotions else 0
        
        summary = f"Recent topics: {'; '.join(recent_topics)}. "
        summary += f"Emotion: {dominant_emotion}. Agent: {most_used_agent}. "
        if voice_percentage > 0:
            summary += f"Voice usage: {voice_percentage:.0f}%. "
        if locations:
            summary += f"Location context: {locations[-1]}."
        
        return summary

    def generate_thread_id(self) -> str:
        """Generate conversation thread ID"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"thread_{timestamp}_{random.randint(1000, 9999)}"

class PremiumAzureVoiceSystem:
    """Premium Azure Voice System with multi-language support"""
    
    def __init__(self):
        self.azure_enabled = AZURE_VOICE_AVAILABLE
        self.basic_voice_enabled = VOICE_AVAILABLE
        self.speech_config = None
        self.speech_recognizer = None
        self.speech_synthesizer = None
        self.tts_engine = None
        self.recognizer = None
        
        if self.azure_enabled:
            try:
                self.setup_azure_voice()
            except Exception as e:
                print(f"Azure Voice setup error: {e}")
                self.azure_enabled = False
        
        if self.basic_voice_enabled:
            try:
                self.setup_basic_voice()
            except Exception as e:
                self.basic_voice_enabled = False
                print(f"Basic voice setup error: {e}")

    def setup_azure_voice(self):
        """Setup premium Azure voice services"""
        azure_key = os.getenv('AZURE_SPEECH_KEY')
        azure_region = os.getenv('AZURE_SPEECH_REGION', 'eastus')
        
        if not azure_key:
            raise Exception("Azure Speech Key not found")
        
        self.speech_config = speechsdk.SpeechConfig(
            subscription=azure_key,
            region=azure_region
        )
        
        self.speech_config.speech_recognition_language = "en-IN"
        self.speech_config.speech_synthesis_voice_name = "en-IN-NeerjaNeural"
        
        audio_config = speechsdk.audio.AudioConfig(use_default_microphone=True)
        self.speech_recognizer = speechsdk.SpeechRecognizer(
            speech_config=self.speech_config,
            audio_config=audio_config
        )
        
        self.speech_synthesizer = speechsdk.SpeechSynthesizer(
            speech_config=self.speech_config
        )

    def setup_basic_voice(self):
        """Setup basic voice as fallback"""
        self.recognizer = sr.Recognizer()
        self.tts_engine = pyttsx3.init()
        
        voices = self.tts_engine.getProperty('voices')
        if voices:
            for voice in voices:
                if 'female' in voice.name.lower() or 'zira' in voice.name.lower():
                    self.tts_engine.setProperty('voice', voice.id)
                    break
        
        self.tts_engine.setProperty('rate', 180)
        self.tts_engine.setProperty('volume', 0.9)

    async def listen_azure(self, language: str = "en-IN") -> Optional[str]:
        """Premium Azure voice recognition"""
        if not self.azure_enabled:
            return None
        
        try:
            self.speech_config.speech_recognition_language = language
            result = self.speech_recognizer.recognize_once()
            
            if result.reason == speechsdk.ResultReason.RecognizedSpeech:
                return result.text
            elif result.reason == speechsdk.ResultReason.NoMatch:
                return None
            elif result.reason == speechsdk.ResultReason.Canceled:
                return None
                
        except Exception as e:
            print(f"Azure Voice error: {e}")
            return None

    def listen_basic(self) -> Optional[str]:
        """Basic voice recognition fallback"""
        if not self.basic_voice_enabled:
            return None
        
        try:
            with sr.Microphone() as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
                audio = self.recognizer.listen(source, timeout=5, phrase_time_limit=15)
            
            return self.recognizer.recognize_google(audio, language='en-IN')
            
        except sr.WaitTimeoutError:
            return None
        except sr.UnknownValueError:
            return None
        except Exception as e:
            print(f"Basic Voice error: {e}")
            return None

    async def listen(self, language: str = "en-IN") -> Optional[str]:
        """Universal listen with Azure priority"""
        if self.azure_enabled:
            result = await self.listen_azure(language)
            if result:
                return result
        
        if self.basic_voice_enabled:
            return self.listen_basic()
        
        return None

    async def speak_azure(self, text: str, language: str = "en-IN", voice: str = None):
        """Premium Azure text-to-speech"""
        if not self.azure_enabled:
            return
        
        try:
            clean_text = self.clean_text_for_tts(text)
            if not clean_text:
                return
            
            voice_map = {
                "en-IN": "en-IN-NeerjaNeural",
                "en-US": "en-US-AriaNeural",
                "hi-IN": "hi-IN-SwaraNeural",
                "en-GB": "en-GB-SoniaNeural"
            }
            
            selected_voice = voice or voice_map.get(language, "en-IN-NeerjaNeural")
            self.speech_config.speech_synthesis_voice_name = selected_voice
            
            result = await asyncio.to_thread(
                self.speech_synthesizer.speak_text, clean_text
            )
            
            if result.reason != speechsdk.ResultReason.SynthesizingAudioCompleted:
                print(f"Azure TTS error: {result.reason}")
                    
        except Exception as e:
            print(f"Azure TTS error: {e}")

    def speak_basic(self, text: str):
        """Basic text-to-speech fallback"""
        if not self.basic_voice_enabled:
            return
        
        try:
            clean_text = self.clean_text_for_tts(text)
            if clean_text:
                self.tts_engine.say(clean_text)
                self.tts_engine.runAndWait()
                
        except Exception as e:
            print(f"Basic TTS error: {e}")

    async def speak(self, text: str, language: str = "en-IN", voice: str = None):
        """Universal speak with Azure priority"""
        if self.azure_enabled:
            await self.speak_azure(text, language, voice)
        elif self.basic_voice_enabled:
            self.speak_basic(text)

    def clean_text_for_tts(self, text: str) -> str:
        """Clean text for professional TTS"""
        if isinstance(text, dict):
            text = text.get('content', str(text))
        
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'#{1,6}\s*(.*?)(?:\n|$)', r'\1', text)
        text = re.sub(r'\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'[🔧💼📈🏥💙🚀🎯📋💡📚🤖⚠️✅❌🔊📝🎤🌤️🌧️⛈️🌙☀️🌍📍💰🔍📰🖼️]', '', text)
        
        if len(text) > 500:
            text = text[:500] + "... Please check the full response on screen."
        
        return text.strip()

    def get_available_voices(self, language: str = "en-IN") -> List[str]:
        """Get available voices for language"""
        voice_options = {
            "en-IN": ["en-IN-NeerjaNeural", "en-IN-PrabhatNeural"],
            "en-US": ["en-US-AriaNeural", "en-US-DavisNeural", "en-US-JennyNeural"],
            "hi-IN": ["hi-IN-SwaraNeural", "hi-IN-MadhurNeural"],
            "en-GB": ["en-GB-SoniaNeural", "en-GB-RyanNeural"]
        }
        
        return voice_options.get(language, ["en-IN-NeerjaNeural"])

class PremiumWebSearchSystem:
    """Premium web search with DuckDuckGo scraping (No API needed)"""
    
    def __init__(self):
        self.search_enabled = True

    def extract_domain(self, url: str) -> str:
        """Extract domain from URL"""
        try:
            from urllib.parse import urlparse
            return urlparse(url).netloc
        except:
            return "Unknown"

    async def search_web(self, query: str, max_results: int = 5) -> Dict[str, Any]:
        """Free DuckDuckGo web search via scraping"""
        try:
            url = f"https://html.duckduckgo.com/html/?q={query}&kl=in-en"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            for result_div in soup.find_all('div', class_='result__body')[:max_results]:
                try:
                    title_elem = result_div.find('a', class_='result__a')
                    snippet_elem = result_div.find('a', class_='result__snippet')
                    
                    if title_elem:
                        title = title_elem.get_text().strip()
                        url_link = title_elem.get('href', '')
                        snippet = snippet_elem.get_text().strip() if snippet_elem else "No description"
                        
                        results.append({
                            "title": title,
                            "snippet": snippet,
                            "url": url_link,
                            "source": self.extract_domain(url_link)
                        })
                except Exception as e:
                    continue
            
            return {
                "success": True,
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            return {"error": f"Web search failed: {e}"}

    async def search_news(self, query: str = "latest news", max_results: int = 5) -> Dict[str, Any]:
        """Free DuckDuckGo news search"""
        try:
            news_query = f"{query} news"
            url = f"https://html.duckduckgo.com/html/?q={news_query}&iar=news&kl=in-en"
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            for result_div in soup.find_all('div', class_='result__body')[:max_results]:
                try:
                    title_elem = result_div.find('a', class_='result__a')
                    snippet_elem = result_div.find('a', class_='result__snippet')
                    
                    if title_elem:
                        title = title_elem.get_text().strip()
                        url_link = title_elem.get('href', '')
                        snippet = snippet_elem.get_text().strip() if snippet_elem else "No description"
                        
                        results.append({
                            "title": title,
                            "snippet": snippet,
                            "url": url_link,
                            "source": self.extract_domain(url_link),
                            "date": "Recent"
                        })
                except Exception as e:
                    continue
            
            return {
                "success": True,
                "query": query,
                "results": results,
                "count": len(results)
            }
            
        except Exception as e:
            return {"error": f"News search failed: {e}"}

class PremiumWeatherSystem:
    """Premium weather system with location detection"""
    
    def __init__(self):
        self.weather_api_key = os.getenv('OPENWEATHER_API_KEY')
        self.weather_enabled = bool(self.weather_api_key)

    async def get_weather(self, location: str = None) -> Dict[str, Any]:
        """Get weather information"""
        if not self.weather_enabled:
            return {"error": "Weather API key not configured"}
        
        try:
            if not location:
                location = self.get_current_location()
            
            url = f"http://api.openweathermap.org/data/2.5/weather"
            params = {
                'q': location,
                'appid': self.weather_api_key,
                'units': 'metric'
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            return {
                "success": True,
                "location": data['name'],
                "country": data['sys']['country'],
                "temperature": data['main']['temp'],
                "feels_like": data['main']['feels_like'],
                "humidity": data['main']['humidity'],
                "pressure": data['main']['pressure'],
                "description": data['weather'][0]['description'],
                "icon": data['weather'][0]['icon'],
                "wind_speed": data['wind']['speed'],
                "visibility": data.get('visibility', 'N/A')
            }
            
        except Exception as e:
            return {"error": f"Weather fetch failed: {e}"}

    def get_current_location(self) -> str:
        """Auto-detect current location"""
        try:
            g = geocoder.ip('me')
            if g.ok:
                return f"{g.city}, {g.country}"
            else:
                return "Mumbai, India"
        except:
            return "Mumbai, India"

    async def get_forecast(self, location: str = None, days: int = 5) -> Dict[str, Any]:
        """Get weather forecast"""
        if not self.weather_enabled:
            return {"error": "Weather API key not configured"}
        
        try:
            if not location:
                location = self.get_current_location()
            
            url = f"http://api.openweathermap.org/data/2.5/forecast"
            params = {
                'q': location,
                'appid': self.weather_api_key,
                'units': 'metric',
                'cnt': days * 8
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            forecast = []
            for item in data['list']:
                forecast.append({
                    "datetime": item['dt_txt'],
                    "temperature": item['main']['temp'],
                    "description": item['weather'][0]['description'],
                    "humidity": item['main']['humidity']
                })
            
            return {
                "success": True,
                "location": data['city']['name'],
                "forecast": forecast
            }
            
        except Exception as e:
            return {"error": f"Forecast fetch failed: {e}"}

class PremiumCryptoSystem:
    """Premium cryptocurrency tracking system"""
    
    def __init__(self):
        self.crypto_enabled = True

    async def get_crypto_prices(self, symbols: List[str] = None) -> Dict[str, Any]:
        """Get cryptocurrency prices"""
        if symbols is None:
            symbols = ['bitcoin', 'ethereum', 'binancecoin', 'cardano', 'solana']
        
        try:
            url = "https://api.coingecko.com/api/v3/simple/price"
            params = {
                'ids': ','.join(symbols),
                'vs_currencies': 'usd,inr',
                'include_24hr_change': 'true',
                'include_market_cap': 'true'
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            prices = {}
            for symbol, info in data.items():
                prices[symbol] = {
                    "usd": info.get('usd', 0),
                    "inr": info.get('inr', 0),
                    "change_24h": info.get('usd_24h_change', 0),
                    "market_cap": info.get('usd_market_cap', 0)
                }
            
            return {
                "success": True,
                "prices": prices,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {"error": f"Crypto price fetch failed: {e}"}

    async def get_trending_cryptos(self) -> Dict[str, Any]:
        """Get trending cryptocurrencies"""
        try:
            url = "https://api.coingecko.com/api/v3/search/trending"
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            data = response.json()
            
            trending = []
            for coin in data['coins']:
                trending.append({
                    "name": coin['item']['name'],
                    "symbol": coin['item']['symbol'],
                    "market_cap_rank": coin['item']['market_cap_rank'],
                    "price_btc": coin['item']['price_btc']
                })
            
            return {
                "success": True,
                "trending": trending
            }
            
        except Exception as e:
            return {"error": f"Trending crypto fetch failed: {e}"}

class PremiumImageGenerationSystem:
    """Premium image generation using Replicate (Free Models)"""
    
    def __init__(self):
        self.replicate_token = os.getenv('REPLICATE_API_TOKEN')
        self.image_gen_enabled = IMAGE_GENERATION_AVAILABLE and bool(self.replicate_token)
        
        if self.image_gen_enabled:
            try:
                replicate.Client(api_token=self.replicate_token)
            except Exception as e:
                print(f"Replicate initialization error: {e}")
                self.image_gen_enabled = False

    async def generate_image(self, prompt: str, size: str = "1024x1024") -> Dict[str, Any]:
        """Generate image using free Replicate SDXL model"""
        if not self.image_gen_enabled:
            return {"error": "Image generation not available"}
        
        try:
            output = await asyncio.to_thread(
                replicate.run,
                "stability-ai/sdxl",
                input={
                    "prompt": prompt,
                    "width": 1024,
                    "height": 1024,
                    "num_outputs": 1,
                    "scheduler": "K_EULER",
                    "num_inference_steps": 20,
                    "guidance_scale": 7.5
                }
            )
            
            image_url = output[0] if output and len(output) > 0 else None
            
            if image_url:
                return {
                    "success": True,
                    "prompt": prompt,
                    "image_url": image_url,
                    "size": size,
                    "model": "SDXL (Replicate)"
                }
            else:
                return {"error": "Replicate returned no image URL"}
                
        except Exception as e:
            return {"error": f"Image generation failed: {e}"}

    async def edit_image(self, image_path: str, prompt: str) -> Dict[str, Any]:
        """Edit image using Replicate (if model supports it)"""
        if not self.image_gen_enabled:
            return {"error": "Image editing not available"}
        
        try:
            return {
                "error": "Image editing not yet implemented with current Replicate models",
                "suggestion": "Try generating a new image with your requirements"
            }
            
        except Exception as e:
            return {"error": f"Image editing failed: {e}"}

class GitHubRepoAnalyzer:
    """Advanced GitHub repository analyzer with bug detection"""
    
    def __init__(self):
        self.active_repo = None
        self.repo_data = {}
        self.qa_engine = None
        self.vector_db_path = None
        
        if GITHUB_INTEGRATION and create_qa_engine:
            try:
                self.qa_engine = create_qa_engine(simple=False)
            except Exception as e:
                try:
                    self.qa_engine = create_qa_engine(simple=True)
                except Exception as e2:
                    print(f"QA Engine initialization failed: {e2}")
                    self.qa_engine = None

    async def analyze_repository(self, repo_url: str) -> Dict[str, Any]:
        """Analyze GitHub repository comprehensively"""
        if not GITHUB_INTEGRATION or not ingest_repo:
            return {"error": "GitHub integration not available"}
        
        try:
            repo_name = repo_url.split('/')[-1].replace('.git', '')
            
            try:
                ingest_repo(repo_url)
            except Exception as e:
                return {"error": f"Failed to ingest repository: {e}"}
            
            self.active_repo = repo_url
            self.repo_data = {
                'name': repo_name,
                'url': repo_url,
                'analyzed_at': datetime.now(),
                'vector_db_path': "./chroma_db"
            }
            
            analysis = await self.perform_code_analysis()
            
            return {
                "success": True,
                "repo_name": repo_name,
                "repo_url": repo_url,
                "analysis": analysis,
                "files_processed": analysis.get('file_count', 0),
                "languages": analysis.get('languages', []),
                "issues_found": analysis.get('issues', []),
                "suggestions": analysis.get('suggestions', [])
            }
            
        except Exception as e:
            return {"error": f"Repository analysis failed: {e}"}

    async def perform_code_analysis(self) -> Dict[str, Any]:
        """Perform comprehensive code analysis"""
        if not self.qa_engine:
            return {
                "error": "QA engine not available",
                'file_count': 'Repository processed',
                'languages': ['Python', 'JavaScript', 'Other'],
                'issues': ["Analysis engine unavailable"],
                'suggestions': ["Manual code review recommended"],
                'detailed_analysis': {}
            }
        
        analysis_questions = [
            "What is the main purpose of this codebase?",
            "What programming languages are used?",
            "Are there any potential bugs or issues in the code?",
            "What improvements can be made to this code?",
            "What is the overall structure and architecture?"
        ]
        
        analysis_results = {}
        for question in analysis_questions:
            try:
                result = self.qa_engine.ask(question)
                if isinstance(result, dict) and 'response' in result:
                    analysis_results[question] = result['response']
                else:
                    analysis_results[question] = str(result)
            except Exception as e:
                analysis_results[question] = f"Analysis failed: {e}"
        
        return {
            'file_count': 'Multiple files processed',
            'languages': ['Python', 'JavaScript', 'Other'],
            'issues': self.extract_issues(analysis_results),
            'suggestions': self.extract_suggestions(analysis_results),
            'detailed_analysis': analysis_results
        }

    def extract_issues(self, analysis: Dict[str, str]) -> List[str]:
        """Extract potential issues from analysis"""
        issues = []
        bug_question = "Are there any potential bugs or issues in the code?"
        
        if bug_question in analysis:
            bug_analysis = analysis[bug_question].lower()
            if 'bug' in bug_analysis or 'issue' in bug_analysis or 'error' in bug_analysis:
                issues.append("Potential bugs detected in codebase")
            if 'security' in bug_analysis:
                issues.append("Security concerns identified")
            if 'performance' in bug_analysis:
                issues.append("Performance optimizations needed")
        
        return issues if issues else ["No critical issues detected"]

    def extract_suggestions(self, analysis: Dict[str, str]) -> List[str]:
        """Extract improvement suggestions"""
        suggestions = []
        improvement_question = "What improvements can be made to this code?"
        
        if improvement_question in analysis:
            suggestions.append("Code structure and architecture improvements")
            suggestions.append("Documentation and comments enhancement")
            suggestions.append("Error handling and validation improvements")
            suggestions.append("Performance optimization opportunities")
        
        return suggestions

    async def answer_repo_question(self, question: str) -> str:
        """Answer questions about the active repository"""
        if not self.active_repo or not self.qa_engine:
            return "No active repository or QA engine not available"
        
        try:
            result = self.qa_engine.ask(question)
            if isinstance(result, dict) and 'response' in result:
                return result['response']
            return str(result)
        except Exception as e:
            return f"Failed to answer repository question: {e}"

    def has_active_repo(self) -> bool:
        """Check if there's an active repository"""
        return self.active_repo is not None

class ProfessionalLanguageDetector:
    """Professional language detection: English + Hinglish only"""
    
    def __init__(self):
        self.language_patterns = {
            "english": [
                "what", "how", "when", "where", "why", "can", "will",
                "should", "could", "would", "the", "and", "but", "this",
                "that", "good", "bad", "right", "wrong", "please", "thank",
                "help", "need", "want", "like", "make", "work", "time"
            ],
            "hinglish": [
                "yaar", "bhai", "dude", "boss", "sir", "madam", "ji",
                "na", "haan", "nahi", "accha", "theek", "kya", "hai",
                "hoon", "main", "tum", "aur", "kar", "kaise", "kyun",
                "matlab", "samjha", "pata", "chal"
            ]
        }
        
        self.hindi_indicators = [
            "hai", "hoon", "kya", "aur", "main", "tum", "yeh", "woh",
            "kaise", "kab", "kahan", "kyun", "matlab", "samjha", "pata"
        ]

    def detect_language(self, text: str) -> str:
        """Detect if text is English or Hinglish"""
        text_words = text.lower().split()
        
        english_count = sum(1 for word in self.language_patterns["english"] if word in text_words)
        hinglish_count = sum(1 for word in self.language_patterns["hinglish"] if word in text_words)
        hindi_count = sum(1 for word in self.hindi_indicators if word in text_words)
        
        total_words = len(text_words)
        if total_words == 0:
            return "english"
        
        if (hinglish_count + hindi_count) / total_words > 0.15:
            return "hinglish"
        
        return "english"

class AdvancedEmotionDetector:
    """PhD-level emotion detection with confidence scoring"""
    
    def __init__(self):
        self.emotion_patterns = {
            "excited": [
                "excited", "amazing", "awesome", "fantastic", "great",
                "wonderful", "thrilled", "happy", "joy", "ecstatic",
                "pumped", "energetic", "love", "brilliant", "perfect"
            ],
            "frustrated": [
                "frustrated", "annoyed", "irritated", "angry", "mad",
                "upset", "pissed", "fed up", "bothered", "stressed",
                "hate", "terrible", "awful", "worst", "stupid"
            ],
            "sad": [
                "sad", "depressed", "down", "blue", "unhappy", "miserable",
                "heartbroken", "grief", "disappointed", "dejected",
                "lonely", "empty", "hopeless", "crying", "tears"
            ],
            "anxious": [
                "anxious", "worried", "nervous", "stressed", "panic",
                "fear", "scared", "concern", "overwhelmed", "tense",
                "pressure", "burden", "difficult", "problem", "issue"
            ],
            "confident": [
                "confident", "sure", "certain", "positive", "optimistic",
                "determined", "ready", "motivated", "strong", "capable",
                "can do", "will do", "believe", "achieve", "success"
            ],
            "confused": [
                "confused", "lost", "unclear", "puzzled", "bewildered",
                "don't understand", "help me understand", "not sure",
                "complicated", "difficult", "hard", "stuck", "doubt"
            ],
            "curious": [
                "curious", "interesting", "wonder", "explore", "discover",
                "learn", "understand", "find out", "know more", "research"
            ]
        }

    def detect_emotion(self, text: str) -> Tuple[str, float]:
        """Detect emotion with confidence score"""
        text_lower = text.lower()
        emotion_scores = {}
        
        for emotion, keywords in self.emotion_patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                emotion_scores[emotion] = score
        
        if not emotion_scores:
            return "neutral", 0.5
        
        detected_emotion = max(emotion_scores, key=emotion_scores.get)
        confidence = emotion_scores[detected_emotion] / len(text_lower.split())
        
        return detected_emotion, min(confidence * 3, 1.0)

class SmartAPIManager:
    """Enhanced API manager with professional providers"""
    
    def __init__(self):
        self.providers = [
            {
                "name": "Groq",
                "url": "https://api.groq.com/openai/v1/chat/completions",
                "models": [
                    "llama-3.1-8b-instant",
                    "llama-3.1-70b-versatile",
                    "mixtral-8x7b-32768"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('GROQ_API_KEY', '')}",
                    "Content-Type": "application/json"
                }
            },
            {
                "name": "OpenRouter",
                "url": "https://openrouter.ai/api/v1/chat/completions",
                "models": [
                    "mistralai/mistral-7b-instruct:free",
                    "meta-llama/llama-3.1-70b-instruct",
                    "anthropic/claude-3.5-sonnet",
                    "google/gemini-pro-1.5"
                ],
                "headers": lambda: {
                    "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY', '')}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://nova-professional.ai",
                    "X-Title": "NOVA Professional"
                }
            }
        ]
        
        self.available = []
        for provider in self.providers:
            key_name = f"{provider['name'].upper()}_API_KEY"
            if os.getenv(key_name):
                self.available.append(provider)
        
        self.current = self.available[0] if self.available else None

    async def get_ai_response(self, user_input: str, system_prompt: str,
                            model_preference: str = None) -> Optional[str]:
        """Get AI response with fallback handling"""
        if not self.current:
            return None
        
        models_to_try = self.current["models"]
        if model_preference:
            preferred_models = [m for m in models_to_try if model_preference in m.lower()]
            models_to_try = preferred_models + [m for m in models_to_try if m not in preferred_models]
        
        for model in models_to_try[:2]:
            try:
                response = requests.post(
                    self.current["url"],
                    headers=self.current["headers"](),
                    json={
                        "model": model,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_input}
                        ],
                        "max_tokens": 1500,
                        "temperature": 0.7,
                        "top_p": 0.9
                    },
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if 'choices' in result and len(result['choices']) > 0:
                        return result['choices'][0]['message']['content'].strip()
                        
            except Exception as e:
                print(f"Model {model} failed: {e}")
                continue
        
        return None

class UltraFileProcessor:
    """Ultra file processing system with all formats"""
    
    def __init__(self):
        self.file_processing_enabled = FILE_PROCESSING_AVAILABLE

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process various file types professionally"""
        if not self.file_processing_enabled:
            return {"error": "File processing not available"}
        
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": "File not found"}
            
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return self.process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.process_word(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self.process_excel(file_path)
            elif file_extension in ['.csv']:
                return self.process_csv(file_path)
            elif file_extension in ['.txt', '.md']:
                return self.process_text(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp']:
                return self.process_image(file_path)
            elif file_extension in ['.py', '.js', '.html', '.css', '.java', '.cpp']:
                return self.process_code(file_path)
            else:
                return {"error": f"Unsupported file type: {file_extension}"}
                
        except Exception as e:
            return {"error": f"File processing failed: {e}"}

    def process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            
            return {
                "file_type": "pdf",
                "pages": len(pdf_reader.pages),
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"PDF processing failed: {e}"}

    def process_word(self, file_path: Path) -> Dict[str, Any]:
        """Process Word files"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "file_type": "word",
                "text": text,
                "word_count": len(text.split()),
                "paragraph_count": len(doc.paragraphs),
                "char_count": len(text),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Word processing failed: {e}"}

    def process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            df = pd.read_excel(file_path)
            
            return {
                "file_type": "excel",
                "shape": df.shape,
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "summary": df.describe().to_dict(),
                "null_values": df.isnull().sum().to_dict(),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Excel processing failed: {e}"}

    def process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            return {
                "file_type": "csv",
                "shape": df.shape,
                "columns": list(df.columns),
                "data_types": df.dtypes.to_dict(),
                "summary": df.describe().to_dict(),
                "sample_data": df.head().to_dict(),
                "null_values": df.isnull().sum().to_dict(),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"CSV processing failed: {e}"}

    def process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            words = text.split()
            lines = text.split('\n')
            
            return {
                "file_type": "text",
                "text": text,
                "word_count": len(words),
                "line_count": len(lines),
                "char_count": len(text),
                "average_words_per_line": len(words) / len(lines) if lines else 0,
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Text processing failed: {e}"}

    def process_image(self, file_path: Path) -> Dict[str, Any]:
        """Process image files"""
        try:
            img = Image.open(file_path)
            
            return {
                "file_type": "image",
                "format": img.format,
                "mode": img.mode,
                "size": img.size,
                "width": img.width,
                "height": img.height,
                "has_transparency": img.mode in ('RGBA', 'LA'),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Image processing failed: {e}"}

    def process_code(self, file_path: Path) -> Dict[str, Any]:
        """Process code files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                code = file.read()
            
            lines = code.split('\n')
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('#')]
            comment_lines = [line for line in lines if line.strip().startswith('#')]
            
            return {
                "file_type": "code",
                "language": file_path.suffix[1:],
                "total_lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "blank_lines": len(lines) - len(code_lines) - len(comment_lines),
                "char_count": len(code),
                "functions": self.count_functions(code, file_path.suffix),
                "success": True
            }
            
        except Exception as e:
            return {"error": f"Code processing failed: {e}"}

    def count_functions(self, code: str, extension: str) -> int:
        """Count functions in code"""
        try:
            if extension == '.py':
                return len(re.findall(r'def\s+\w+', code))
            elif extension == '.js':
                return len(re.findall(r'function\s+\w+', code))
            elif extension in ['.java', '.cpp', '.c']:
                return len(re.findall(r'\w+\s+\w+\s*\([^)]*\)\s*{', code))
            else:
                return 0
        except:
            return 0

# ========== NOVA ULTRA PROFESSIONAL BACKEND ==========
class NovaUltraProfessionalBackend:
    """NOVA Ultra Professional Backend - Complete Integrated System with ALL features"""
    
    def __init__(self):
        """Initialize the complete ultra professional system"""
        # Core systems
        self.memory = UltraHybridMemorySystem()
        self.language_detector = ProfessionalLanguageDetector()
        self.emotion_detector = AdvancedEmotionDetector()
        self.api_manager = SmartAPIManager()
        
        # Premium systems
        self.voice_system = PremiumAzureVoiceSystem()
        self.web_search = PremiumWebSearchSystem()
        self.weather_system = PremiumWeatherSystem()
        self.crypto_system = PremiumCryptoSystem()
        self.image_generator = PremiumImageGenerationSystem()
        self.file_processor = UltraFileProcessor()
        self.github_analyzer = GitHubRepoAnalyzer()
        
        # Load professional agents
        self.agents = {}
        if PROFESSIONAL_AGENTS_LOADED:
            try:
                self.agents = {
                    'coding': ProLevelCodingExpert(),
                    'career': ProfessionalCareerCoach(),
                    'business': SmartBusinessConsultant(),
                    'medical': SimpleMedicalAdvisor(),
                    'emotional': SimpleEmotionalCounselor(),
                    'technical_architect': TechnicalArchitect()
                }
            except Exception as e:
                print(f"Agent loading error: {e}")
        
        # Advanced systems
        self.orchestrator = None
        self.drift_detector = None
        if ADVANCED_SYSTEMS:
            try:
                self.smart_orchestrator = IntelligentAPIOrchestrator()
                self.api_drift_detector = APIPerformanceDrifter()
            except Exception as e:
                print(f"Advanced systems error: {e}")
        
        # Session management
        self.current_session = f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        self.user_id = "ultra_professional_user"
        self.conversation_count = 0
        self.current_agent = "general"
        self.current_location = None
        
        # Premium agent patterns (enhanced)
        self.agent_patterns = {
            "coding": {
                "keywords": [
                    "code", "programming", "debug", "error", "python", "javascript",
                    "html", "css", "react", "nodejs", "api", "database", "software",
                    "development", "algorithm", "function", "bug", "git", "github",
                    "frontend", "backend", "fullstack", "deployment", "testing"
                ],
                "system_prompt": "You are NOVA Coding Expert, a senior software engineer with expertise in all programming languages and frameworks. Provide production-ready solutions with best practices, error handling, and optimization."
            },
            "career": {
                "keywords": [
                    "resume", "interview", "job", "career", "hiring", "cv", "cover letter",
                    "application", "salary", "promotion", "skills", "experience",
                    "professional", "employment", "linkedin", "networking", "portfolio"
                ],
                "system_prompt": "You are NOVA Career Coach, a professional career advisor with 20+ years experience. Provide expert guidance on resumes, interviews, career planning, and professional development."
            },
            "business": {
                "keywords": [
                    "business", "analysis", "kpi", "metrics", "roi", "revenue", "profit",
                    "growth", "strategy", "market", "data", "analytics", "reporting",
                    "process", "optimization", "dashboard", "startup", "investment"
                ],
                "system_prompt": "You are NOVA Business Consultant, an expert business analyst with deep knowledge of KPIs, metrics, and business intelligence. Provide strategic insights and data-driven recommendations."
            },
            "medical": {
                "keywords": [
                    "health", "doctor", "medicine", "symptoms", "illness", "pain",
                    "fever", "medical", "treatment", "diagnosis", "medication",
                    "healthcare", "wellness", "fitness", "nutrition", "mental health"
                ],
                "system_prompt": "You are Dr. NOVA, a medical expert with comprehensive knowledge. Provide medical insights while always emphasizing the importance of consulting healthcare professionals."
            },
            "emotional": {
                "keywords": [
                    "sad", "stress", "anxiety", "depression", "emotional", "feelings",
                    "therapy", "mental", "mood", "worried", "upset", "frustrated",
                    "lonely", "overwhelmed", "counseling", "support"
                ],
                "system_prompt": "You are Dr. NOVA Counselor, a compassionate therapist with PhD-level emotional intelligence. Provide empathetic support and practical guidance."
            },
            "technical_architect": {
                "keywords": [
                    "architecture", "system design", "microservice", "monolith",
                    "serverless", "event driven", "scalability", "throughput",
                    "design pattern", "distributed", "load balancer", "database design",
                    "high availability", "fault tolerance", "performance", "scaling"
                ],
                "system_prompt": "You are NOVA Technical Architect, a senior system designer with expertise in distributed systems, scalability, and architectural patterns. Provide comprehensive architectural guidance with best practices."
            },
            "search": {
                "keywords": [
                    "search", "find", "look up", "google", "information", "research",
                    "web", "internet", "news", "latest", "current", "trending"
                ],
                "system_prompt": "You are NOVA Search Assistant, helping users find accurate and relevant information from the web. Provide comprehensive search results with reliable sources."
            },
            "weather": {
                "keywords": [
                    "weather", "temperature", "rain", "sunny", "cloudy", "forecast",
                    "climate", "humidity", "wind", "storm", "snow", "hot", "cold"
                ],
                "system_prompt": "You are NOVA Weather Assistant, providing accurate weather information and forecasts. Help users plan their activities based on weather conditions."
            },
            "crypto": {
                "keywords": [
                    "crypto", "cryptocurrency", "bitcoin", "ethereum", "trading",
                    "blockchain", "price", "market", "investment", "coin", "defi"
                ],
                "system_prompt": "You are NOVA Crypto Assistant, providing cryptocurrency market insights and information. Always mention investment risks and suggest professional financial advice."
            },
            "image": {
                "keywords": [
                    "image", "picture", "generate", "create", "draw", "art", "design",
                    "visual", "illustration", "photo", "graphic", "artwork"
                ],
                "system_prompt": "You are NOVA Image Assistant, helping users create and understand images. Provide creative guidance for image generation and visual content."
            }
        }

    async def detect_agent_type(self, user_input: str, context: str = "") -> Tuple[str, float]:
        """Detect appropriate agent with ultra confidence scoring"""
        text_lower = (user_input + " " + context).lower()
        agent_scores = {}
        
        for agent_name, agent_data in self.agent_patterns.items():
            keywords = agent_data["keywords"]
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                agent_scores[agent_name] = score
        
        if not agent_scores:
            return "general", 0.0
        
        best_agent = max(agent_scores, key=agent_scores.get)
        confidence = agent_scores[best_agent] / len(text_lower.split())
        
        return best_agent, min(confidence, 1.0)

    async def create_ultra_professional_prompt(self, agent_type: str, language: str,
                                             emotion: str, context: str = "",
                                             location: str = None, weather: str = None) -> str:
        """Create ultra professional system prompt"""
        base_personality = """You are NOVA Ultra Professional AI, the most advanced assistant with expertise across ALL domains. You provide:

- Ultra professional-grade assistance with enterprise-level quality
- PhD-level expertise and emotional intelligence across all fields
- Context-aware responses with perfect memory and location awareness
- Practical, actionable solutions with detailed explanations
- Multi-domain knowledge spanning technology, business, health, career, weather, crypto, and more
- Real-time information access and premium features

You are NOT a robotic AI - you understand emotions, context, location, and provide genuinely helpful guidance with ultra-premium quality."""
        
        agent_data = self.agent_patterns.get(agent_type, {})
        agent_prompt = agent_data.get("system_prompt", "Provide general assistance with ultra professional quality.")
        
        if language == "hinglish":
            language_instruction = "Mix Hindi और English naturally when appropriate, maintaining ultra professional tone. Use Hinglish only when user prefers it."
        else:
            language_instruction = "Communicate in clear, ultra professional English with appropriate technical depth."
        
        emotion_adaptations = {
            "excited": "Match user's enthusiasm while providing comprehensive guidance to channel their energy productively with premium quality.",
            "frustrated": "Provide patient, solution-focused assistance to resolve their concerns with empathy and clear step-by-step guidance.",
            "sad": "Offer compassionate support with gentle encouragement and practical guidance to improve their situation with care.",
            "anxious": "Provide calm, reassuring guidance with step-by-step solutions to reduce anxiety and build confidence.",
            "confident": "Support their confidence with detailed information and advanced guidance to maintain momentum.",
            "confused": "Provide crystal clear, step-by-step explanations with examples to eliminate confusion completely.",
            "curious": "Feed their curiosity with comprehensive, detailed information and encourage exploration."
        }
        
        emotion_instruction = emotion_adaptations.get(emotion, "Respond with appropriate emotional intelligence and ultra professional empathy.")
        
        context_instruction = ""
        if context:
            context_instruction = f"\n\nCONVERSATION CONTEXT:\n{context}\n\nUse this context to provide personalized, contextually aware responses."
        
        location_instruction = ""
        if location:
            location_instruction = f"\n\nUSER LOCATION: {location}"
            if weather:
                location_instruction += f"\nCURRENT WEATHER: {weather}"
            location_instruction += "\nConsider location and weather in your responses when relevant."
        
        return f"""{base_personality}

SPECIALIST MODE: {agent_prompt}

COMMUNICATION: {language_instruction}

EMOTIONAL INTELLIGENCE: {emotion_instruction}

{context_instruction}

{location_instruction}

Provide ultra comprehensive, accurate, and professionally helpful responses that demonstrate expertise and genuine care for the user's needs with premium quality."""

    PLACEHOLDERS = {
        "consultation response", "detailed explanation", "concept explanation",
        "career guidance response", "business consultation response",
        "medical guidance response", "emotional support response", 
        "architecture solution", ""
    }
    
    def _is_placeholder(self, text: str) -> bool:
        """Check if text is a placeholder"""
        return not text or text.strip().lower() in self.PLACEHOLDERS

    async def get_ultra_professional_response(self, user_input: str, user_id: str = "default_user", session_id: str = None) -> Dict[str, Any]:
        """Get ultra comprehensive professional response"""
        start_time = datetime.now()
        
        try:
            context = self.memory.get_relevant_context(user_input, user_id, 15)
            
            language = self.language_detector.detect_language(user_input)
            emotion, emotion_confidence = self.emotion_detector.detect_emotion(user_input)
            agent_type, agent_confidence = await self.detect_agent_type(user_input, context)
            
            if not self.current_location:
                self.current_location = self.weather_system.get_current_location()
            
            weather_context = None
            if any(keyword in user_input.lower() for keyword in ['weather', 'temperature', 'rain', 'sunny']):
                weather_result = await self.weather_system.get_weather(self.current_location)
                if weather_result.get('success'):
                    weather_context = f"{weather_result['temperature']}°C, {weather_result['description']}"
            
            system_prompt = await self.create_ultra_professional_prompt(
                agent_type, language, emotion, context, self.current_location, weather_context
            )
            
            response = None
            if agent_type in self.agents and PROFESSIONAL_AGENTS_LOADED:
                try:
                    agent = self.agents[agent_type]
                    if agent_type == 'coding':
                        result = await agent.understand_and_solve(user_input, context)
                    elif agent_type == 'career':
                        result = await agent.provide_career_guidance(user_input)
                    elif agent_type == 'business':
                        result = await agent.provide_business_consultation(user_input)
                    elif agent_type == 'medical':
                        result = await agent.provide_health_guidance(user_input)
                    elif agent_type == 'emotional':
                        result = await agent.provide_support(user_input)
                    elif agent_type == 'technical_architect':
                        result = await agent.provide_architecture_guidance(user_input, context)
                    
                    if isinstance(result, dict):
                        candidate = (result.get("content") or 
                                   result.get("response") or 
                                   result.get("answer") or
                                   result.get("message"))
                    elif isinstance(result, str):
                        candidate = result
                    else:
                        candidate = None
                    
                    if candidate and not self._is_placeholder(candidate):
                        response = candidate
                    else:
                        print(f"Agent placeholder detected ('{candidate}') → fallback to LLM")
                        
                except Exception as e:
                    print(f"Agent {agent_type} error: {e}")
            
            if not response:
                model_preference = {
                    'coding': 'llama',
                    'business': 'llama',
                    'medical': 'llama',
                    'career': 'llama',
                    'search': 'llama',
                    'weather': 'llama',
                    'crypto': 'llama',
                    'image': 'llama'
                }.get(agent_type, 'llama')
                
                response = await self.api_manager.get_ai_response(
                    user_input, system_prompt, model_preference
                )
            
            if not response:
                response = self.get_ultra_professional_fallback(user_input, agent_type, language, emotion)
            
            response_time = (datetime.now() - start_time).total_seconds()
            
            await self.memory.remember_conversation(
                user_id, session_id or self.current_session, user_input, response,
                agent_type, language, emotion, agent_confidence,
                intent="general", response_time=response_time,
                voice_used=False, location=self.current_location,
                weather_context=weather_context
            )
            
            return {
                "success": True,
                "response": response,
                "agent_type": agent_type,
                "agent_confidence": agent_confidence,
                "language": language,
                "emotion": emotion,
                "emotion_confidence": emotion_confidence,
                "response_time": response_time,
                "context_used": bool(context),
                "location": self.current_location,
                "weather_context": weather_context
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "response": f"I encountered an error processing your request: {e}",
                "agent_type": "error",
                "response_time": (datetime.now() - start_time).total_seconds()
            }

    def get_ultra_professional_fallback(self, user_input: str, agent_type: str,
                                      language: str, emotion: str) -> str:
        """Ultra professional fallback responses"""
        fallback_responses = {
            "coding": "I understand you need coding assistance. While I'm experiencing technical difficulties with my advanced systems, I can help with general programming guidance. Please describe your specific coding challenge, and I'll provide the best assistance possible.",
            "career": "I'm here to help with your career development needs. Even with limited connectivity, I can provide guidance on resume building, interview preparation, and career planning. What specific career area would you like to focus on?",
            "business": "I can assist with business analysis and strategic insights. While my advanced business intelligence tools are temporarily unavailable, I can help with general business guidance, process optimization, and strategic thinking. What business challenge can I help you with?",
            "medical": "I understand you have health-related questions. While I'm experiencing connectivity issues with my medical database, I can provide general health information. Please remember that for serious health concerns, it's important to consult with healthcare professionals.",
            "emotional": "I'm here to provide emotional support and guidance. Even when my systems are limited, I genuinely care about your wellbeing. Please share what's on your mind, and I'll do my best to provide helpful support and coping strategies.",
            "search": "I can help you find information. While my web search capabilities are temporarily limited, I can provide general guidance based on my knowledge. What would you like to know?",
            "weather": "I can provide weather assistance. While my real-time weather services are temporarily unavailable, I can offer general weather guidance and planning advice.",
            "crypto": "I can help with cryptocurrency information. While my live market data is temporarily unavailable, I can provide general crypto guidance and educational information.",
            "image": "I can assist with image-related tasks. While my image generation services are temporarily limited, I can provide guidance on visual design and creativity."
        }
        
        base_response = fallback_responses.get(agent_type,
            "I'm experiencing some technical difficulties but I'm still here to help you. Please let me know what you need assistance with, and I'll provide the best guidance I can.")
        
        if language == "hinglish":
            base_response += "\n\nAap Hindi mein bhi puch sakte hain - main dono languages mein help kar sakta hun!"
        
        emotion_additions = {
            "frustrated": "\n\nI understand this might be frustrating. Let me focus on providing you with practical solutions to resolve your concerns.",
            "sad": "\n\nI can sense you might be going through a difficult time. I'm here to provide support and guidance to help improve your situation.",
            "anxious": "\n\nI understand you might be feeling anxious about this. Let me provide clear, step-by-step guidance to help address your concerns.",
            "excited": "\n\nI can feel your enthusiasm! Let me help channel that energy into getting you the comprehensive assistance you need.",
            "curious": "\n\nI love your curiosity! Let me provide detailed information to satisfy your quest for knowledge."
        }
        
        if emotion in emotion_additions:
            base_response += emotion_additions[emotion]
        
        return base_response

    async def handle_voice_input(self, audio_data: str, language: str = "en-IN", user_id: str = "default_user") -> Dict[str, Any]:
        """Handle voice input and return response"""
        if not self.voice_system.azure_enabled and not self.voice_system.basic_voice_enabled:
            return {"error": "Voice system not available"}
        
        try:
            voice_engine = "Azure" if self.voice_system.azure_enabled else "Basic"
            
            # Process audio data (in a real implementation, you would decode the base64 and process it)
            user_input = "Sample voice input"  # Replace with actual voice recognition
            
            result = await self.get_ultra_professional_response(user_input, user_id)
            
            self.memory.remember_voice_interaction(
                user_id, user_input, result.get('response', ''),
                result.get('language', 'english'), result.get('emotion', 'neutral'),
                voice_engine
            )
            
            return result
            
        except Exception as e:
            return {"error": f"Voice processing failed: {e}"}

    async def process_file(self, file_path: str, user_id: str = "default_user") -> Dict[str, Any]:
        """Process file and return results"""
        result = self.file_processor.process_file(file_path)
        
        if 'error' not in result:
            self.memory.remember_file_processing(
                user_id, file_path, result.get('file_type', 'unknown'),
                str(result), True
            )
        else:
            self.memory.remember_file_processing(
                user_id, file_path, "unknown", result['error'], False
            )
        
        return result

    async def analyze_github_repo(self, repo_url: str, user_id: str = "default_user") -> Dict[str, Any]:
        """Analyze GitHub repository"""
        result = await self.github_analyzer.analyze_repository(repo_url)
        return result

    async def answer_repo_question(self, question: str, user_id: str = "default_user") -> Dict[str, Any]:
        """Answer question about active repository"""
        if not self.github_analyzer.has_active_repo():
            return {"error": "No active repository"}
        
        response = await self.github_analyzer.answer_repo_question(question)
        
        return {
            "success": True,
            "response": response,
            "agent_type": "github_analyzer",
            "language": "english",
            "emotion": "neutral"
        }

    async def get_weather(self, location: str = None, user_id: str = "default_user") -> Dict[str, Any]:
        """Get weather information"""
        result = await self.weather_system.get_weather(location)
        return result

    async def get_weather_forecast(self, location: str = None, days: int = 5, user_id: str = "default_user") -> Dict[str, Any]:
        """Get weather forecast"""
        result = await self.weather_system.get_forecast(location, days)
        return result

    async def search_web(self, query: str, max_results: int = 5, user_id: str = "default_user") -> Dict[str, Any]:
        """Search the web"""
        result = await self.web_search.search_web(query, max_results)
        if result.get('success'):
            self.memory.remember_search_query(user_id, query, "web", len(result['results']))
        return result

    async def search_news(self, query: str = "latest news", max_results: int = 5, user_id: str = "default_user") -> Dict[str, Any]:
        """Search news"""
        result = await self.web_search.search_news(query, max_results)
        if result.get('success'):
            self.memory.remember_search_query(user_id, query, "news", len(result['results']))
        return result

    async def get_crypto_prices(self, symbols: List[str] = None, user_id: str = "default_user") -> Dict[str, Any]:
        """Get cryptocurrency prices"""
        result = await self.crypto_system.get_crypto_prices(symbols)
        return result

    async def get_trending_cryptos(self, user_id: str = "default_user") -> Dict[str, Any]:
        """Get trending cryptocurrencies"""
        result = await self.crypto_system.get_trending_cryptos()
        return result

    async def generate_image(self, prompt: str, size: str = "1024x1024", user_id: str = "default_user") -> Dict[str, Any]:
        """Generate image"""
        result = await self.image_generator.generate_image(prompt, size)
        if result.get('success'):
            self.memory.image_memory.append({
                'prompt': prompt,
                'url': result['image_url'],
                'timestamp': datetime.now()
            })
        return result

    async def get_conversation_history(self, user_id: str, limit: int = 15) -> Dict[str, Any]:
        """Get conversation history"""
        try:
            with sqlite3.connect(self.memory.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute('''
                SELECT user_input, bot_response, agent_type, emotion, timestamp
                FROM conversations 
                WHERE user_id = ? 
                ORDER BY timestamp DESC 
                LIMIT ?
                ''', (user_id, limit))
                conversations = cursor.fetchall()
            
            return {
                "success": True,
                "conversations": [
                    {
                        "user_input": conv[0],
                        "bot_response": conv[1],
                        "agent_type": conv[2],
                        "emotion": conv[3],
                        "timestamp": conv[4]
                    } for conv in conversations
                ]
            }
        except Exception as e:
            return {"error": f"Failed to get conversation history: {e}"}

    async def get_system_status(self) -> Dict[str, Any]:
        """Get system status"""
        return {
            "professional_agents_loaded": PROFESSIONAL_AGENTS_LOADED,
            "azure_voice_available": AZURE_VOICE_AVAILABLE,
            "github_integration": GITHUB_INTEGRATION,
            "web_search_available": WEB_SEARCH_AVAILABLE,
            "weather_system_enabled": self.weather_system.weather_enabled,
            "crypto_system_enabled": self.crypto_system.crypto_enabled,
            "image_generation_enabled": self.image_generator.image_gen_enabled,
            "file_processing_available": FILE_PROCESSING_AVAILABLE,
            "current_location": self.current_location,
            "active_repository": self.github_analyzer.active_repo if self.github_analyzer.has_active_repo() else None,
            "voice_engine": "Azure Premium" if self.voice_system.azure_enabled else "Basic" if self.voice_system.basic_voice_enabled else "None"
        }

# ========== FASTAPI ROUTES ==========
nova_backend = NovaUltraProfessionalBackend()

@app.get("/")
async def root():
    return {"message": "NOVA Ultra Professional AI System is running"}

@app.post("/api/chat")
async def chat_endpoint(input_data: UserInput):
    """Main chat endpoint"""
    try:
        result = await nova_backend.get_ultra_professional_response(
            input_data.text, 
            input_data.user_id,
            input_data.session_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/voice")
async def voice_endpoint(voice_input: VoiceInput):
    """Voice input endpoint"""
    try:
        result = await nova_backend.handle_voice_input(
            voice_input.audio_data,
            voice_input.language,
            voice_input.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/file/process")
async def process_file_endpoint(file_request: FileProcessRequest):
    """File processing endpoint"""
    try:
        result = await nova_backend.process_file(
            file_request.file_path,
            file_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/github/analyze")
async def analyze_github_endpoint(github_request: GitHubAnalysisRequest):
    """GitHub analysis endpoint"""
    try:
        result = await nova_backend.analyze_github_repo(
            github_request.repo_url,
            github_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/github/ask")
async def ask_github_endpoint(input_data: UserInput):
    """GitHub question endpoint"""
    try:
        result = await nova_backend.answer_repo_question(
            input_data.text,
            input_data.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/weather")
async def weather_endpoint(weather_request: WeatherRequest):
    """Weather endpoint"""
    try:
        result = await nova_backend.get_weather(
            weather_request.location,
            weather_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/weather/forecast")
async def weather_forecast_endpoint(location: str = None, days: int = 5, user_id: str = "default_user"):
    """Weather forecast endpoint"""
    try:
        result = await nova_backend.get_weather_forecast(
            location,
            days,
            user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/search/web")
async def web_search_endpoint(search_request: SearchRequest):
    """Web search endpoint"""
    try:
        result = await nova_backend.search_web(
            search_request.query,
            search_request.max_results,
            search_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/search/news")
async def news_search_endpoint(search_request: SearchRequest):
    """News search endpoint"""
    try:
        result = await nova_backend.search_news(
            search_request.query,
            search_request.max_results,
            search_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/crypto/prices")
async def crypto_prices_endpoint(crypto_request: CryptoRequest):
    """Crypto prices endpoint"""
    try:
        result = await nova_backend.get_crypto_prices(
            crypto_request.symbols,
            crypto_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/crypto/trending")
async def crypto_trending_endpoint(user_id: str = "default_user"):
    """Trending cryptos endpoint"""
    try:
        result = await nova_backend.get_trending_cryptos(user_id)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/image/generate")
async def image_generate_endpoint(image_request: ImageGenerationRequest):
    """Image generation endpoint"""
    try:
        result = await nova_backend.generate_image(
            image_request.prompt,
            image_request.size,
            image_request.user_id
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/memory/conversations")
async def get_conversation_history(user_id: str, limit: int = 15):
    """Get conversation history"""
    try:
        result = await nova_backend.get_conversation_history(user_id, limit)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/status")
async def get_status():
    """Get system status"""
    try:
        result = await nova_backend.get_system_status()
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ========== MAIN EXECUTION ==========
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)